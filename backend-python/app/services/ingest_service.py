from app.db.models import Document
from app.db.session import async_session
from app.services.cache_service import cache_service
from app.utils.cached_embeddings import create_cached_embeddings
# from langchain.embeddings import OpenAIEmbeddings, HuggingFaceEmbeddings
# from langchain.text_splitter import TextSplitter
# from langchain.vectorstores.pgvector import PGVector

import os
from tempfile import NamedTemporaryFile
from langchain_openai import OpenAIEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores.pgvector import PGVector
from sqlalchemy.ext.asyncio import AsyncSession
from app.db.models import Document
from app.db.session import async_session

async def ingest_document(file):
    # Save uploaded file to temp location
    suffix = os.path.splitext(file.filename)[1]
    with NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    # Parse file content
    if suffix == ".txt":
        with open(tmp_path, "r", encoding="utf-8") as f:
            text = f.read() 
    elif suffix == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(tmp_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(tmp_path)
        text = "\n".join([p.text for p in doc.paragraphs])
    else:
        os.remove(tmp_path)
        return {"status": "error", "message": "Unsupported file type"}

    # Generate embeddings with caching (choose model)
    # base_embeddings = OpenAIEmbeddings()  # or HuggingFaceEmbeddings()
    base_embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    embeddings = create_cached_embeddings(base_embeddings, "sentence-transformers/all-MiniLM-L6-v2")

    # Split text
    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    docs = splitter.split_text(text)

    # Store in pgvector via Langchain
    vectorstore = PGVector(
        connection_string=os.getenv("PGVECTOR_CONN", "postgresql+psycopg2://postgres:postgres@localhost:5432/postgres"),
        embedding_function=embeddings,
        collection_name="documents"
    )

    # Store metadata in DB and get document id
    async with async_session() as session:
        async with session.begin():
            doc = Document(filename=file.filename, num_chunks=len(docs), vector_ids=[])
            session.add(doc)
            await session.flush()  # get doc.id
            document_id = str(doc.id)

            # Add document_id metadata to each chunk
            metadatas = [{"document_id": document_id} for _ in docs]
            ids = vectorstore.add_texts(docs, metadatas=metadatas)
            doc.vector_ids = ids

    # Invalidate any cached queries since we've added new content
    # Note: We can't invalidate specific queries, so we'll rely on TTL
    # In a production system, you might want to implement more sophisticated cache invalidation
    
    os.remove(tmp_path)
    return {"status": "success", "filename": file.filename, "chunks": len(docs)}
